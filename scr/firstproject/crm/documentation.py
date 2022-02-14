from docx import Document
from docx.shared import RGBColor, Pt
from docx.text.run import Run
from typing import Dict


class DocumentPreparation:

    def __init__(self, document_input_path: str,
                 document_output_path: str,
                 data: Dict):
        self.doc = Document(document_input_path)
        self.output_path = document_output_path
        self.data = data

    def helper_method(self, p):
        for run in p.runs:
            label = run.text[1:-1]
            if label in self.data:
                sample = str(self.data[label]).split(' ')[::-1]
                for string in sample:
                    new_run_element = p._element._new_r()
                    run._element.addnext(new_run_element)
                    new_run = Run(new_run_element, run._parent)
                    new_run.text = string
                    new_run.bold = True
                    new_run.font.color.rgb = RGBColor(255, 0, 0)
                    new_run.font.size = Pt(10)

                    new_run_element = p._element._new_r()
                    run._element.addnext(new_run_element)
                    new_run = Run(new_run_element, run._parent)
                    new_run.text = ' '
                    new_run.bold = True
                    new_run.font.color.rgb = RGBColor(255, 0, 0)
                    new_run.font.size = Pt(10)
                run.clear()

    def fill_paragraphs_with_data(self):
        for p in self.doc.paragraphs:
            self.helper_method(p)

    def fill_tables_with_data(self):

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.helper_method(p)

    def generate_document(self):
        self.fill_paragraphs_with_data()
        self.fill_tables_with_data()
        self.doc.save(self.output_path)
        # return self.doc
